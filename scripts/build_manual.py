"""
Builds dist/MetaVibing-Field-Manual-v0.1.pdf and .docx from book/manuscript.md.

Regenerate with:
    python scripts/build_manual.py

Requires: reportlab, python-docx (both already in this environment).
Regenerate book/assets/ first with scripts/gen_assets.py if the diagrams
or logo change.
"""
import re
import pathlib
import sys

sys.path.insert(0, str(pathlib.Path(__file__).parent))
from md_parse import parse, inline_to_reportlab, strip_inline  # noqa: E402

ROOT = pathlib.Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "book" / "manuscript.md"
ASSETS = ROOT / "book" / "assets"
DIST = ROOT / "dist"
DIST.mkdir(exist_ok=True)

VERSION_LABEL = "v0.1 Research Preview"
TITLE = "MetaVibing"
SUBTITLE = "A Field Manual for Evolving Your AI Collaborator"

INDIGO = "#4F46E5"
TEAL = "#0D9488"
VIOLET = "#7C3AED"
AMBER = "#D97706"
DARK = "#1F2937"
GRAY = "#6B7280"
LIGHT = "#F3F4F6"

DIAGRAM_FILES = {
    "core-loop": ASSETS / "diagram-core-loop.png",
    "meta-stack": ASSETS / "diagram-meta-stack.png",
    "three-strikes": ASSETS / "diagram-three-strikes.png",
    "friction-artifact": ASSETS / "diagram-friction-artifact.png",
}

CALLOUT_COLORS = {
    "principle": INDIGO,
    "practice": TEAL,
    "warning": AMBER,
    "evidence": VIOLET,
}

PART_RE = re.compile(r"^Part\s+[IVXLC]+\b")


def load_blocks():
    text = MANUSCRIPT.read_text(encoding="utf-8")
    blocks = parse(text)
    assert blocks[0][0] == "h1", "manuscript must start with a single # title"
    return blocks[1:]  # drop the h1 -- cover is built separately


# ═════════════════════════════════════════════════════════════════════════
# PDF (reportlab)
# ═════════════════════════════════════════════════════════════════════════

def build_pdf(blocks, out_path):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.platypus import (
        BaseDocTemplate, PageTemplate, Frame, Paragraph, Spacer, PageBreak,
        Table, TableStyle, Image as RLImage, Preformatted, KeepTogether,
        NextPageTemplate, FrameBreak,
    )

    PAGE_W, PAGE_H = LETTER
    MARGIN = 0.9 * inch

    styles = getSampleStyleSheet()

    def pstyle(name, **kw):
        base = dict(fontName="Helvetica", fontSize=10.5, leading=15,
                    textColor=colors.HexColor(DARK), spaceAfter=8)
        base.update(kw)
        return ParagraphStyle(name, **base)

    body_style = pstyle("Body")
    h2_style = pstyle("H2", fontName="Helvetica-Bold", fontSize=16, leading=20,
                       textColor=colors.HexColor(INDIGO), spaceBefore=18, spaceAfter=10)
    h3_style = pstyle("H3", fontName="Helvetica-Bold", fontSize=13, leading=17,
                       textColor=colors.HexColor(TEAL), spaceBefore=14, spaceAfter=8)
    h4_style = pstyle("H4", fontName="Helvetica-Bold", fontSize=11.5, leading=15,
                       textColor=colors.HexColor(DARK), spaceBefore=10, spaceAfter=6)
    part_style = pstyle("Part", fontName="Helvetica-Bold", fontSize=30, leading=36,
                         textColor=colors.HexColor(INDIGO), alignment=TA_CENTER)
    quote_style = pstyle("Quote", fontName="Helvetica-Oblique", fontSize=11.5, leading=16,
                          textColor=colors.HexColor(DARK), leftIndent=24, spaceBefore=6, spaceAfter=10)
    li_style = pstyle("LI", leftIndent=18, bulletIndent=6, spaceAfter=4)
    callout_label_style = pstyle("CalloutLabel", fontName="Helvetica-Bold", fontSize=9.5,
                                  spaceAfter=4, textColor=colors.white)
    callout_body_style = pstyle("CalloutBody", spaceAfter=4)
    caption_style = pstyle("Caption", fontName="Helvetica-Oblique", fontSize=9,
                            textColor=colors.HexColor(GRAY), alignment=TA_CENTER, spaceBefore=4, spaceAfter=14)
    code_style = ParagraphStyle("Code", fontName="Courier", fontSize=8.6, leading=11.5,
                                 textColor=colors.HexColor(DARK), backColor=colors.HexColor(LIGHT),
                                 borderPadding=8, spaceBefore=6, spaceAfter=10)

    story = []

    # ── Cover ──────────────────────────────────────────────────────────────
    story.append(Spacer(1, 1.6 * inch))
    logo_path = ASSETS / "logo.png"
    if logo_path.exists():
        img = RLImage(str(logo_path), width=1.4 * inch, height=1.4 * inch)
        img.hAlign = "CENTER"
        story.append(img)
    story.append(Spacer(1, 0.5 * inch))
    story.append(Paragraph(TITLE, pstyle("CoverTitle", fontName="Helvetica-Bold",
                                          fontSize=38, leading=46, alignment=TA_CENTER,
                                          textColor=colors.HexColor(INDIGO))))
    story.append(Spacer(1, 0.25 * inch))
    story.append(Paragraph(SUBTITLE, pstyle("CoverSub", fontName="Helvetica",
                                             fontSize=15, leading=20, alignment=TA_CENTER,
                                             textColor=colors.HexColor(DARK))))
    story.append(Spacer(1, 0.4 * inch))
    story.append(Paragraph(VERSION_LABEL, pstyle("CoverVersion", fontName="Helvetica-Bold",
                                                  fontSize=11, leading=14, alignment=TA_CENTER,
                                                  textColor=colors.HexColor(TEAL))))
    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())

    # ── Body ───────────────────────────────────────────────────────────────

    def render_inline(text):
        return inline_to_reportlab(text)

    def render_blocks_into(blocks, target, in_callout=False):
        i = 0
        n = len(blocks)
        while i < n:
            b = blocks[i]
            kind = b[0]

            if kind == "h2":
                text = b[1]
                if PART_RE.match(text):
                    target.append(PageBreak())
                    target.append(Spacer(1, 2.2 * inch))
                    target.append(Paragraph(render_inline(text), part_style))
                    target.append(PageBreak())
                else:
                    target.append(Paragraph(render_inline(text), h2_style))
            elif kind == "h3":
                target.append(Paragraph(render_inline(b[1]), h3_style))
            elif kind in ("h4", "h5", "h6"):
                target.append(Paragraph(render_inline(b[1]), h4_style))
            elif kind == "p":
                style = callout_body_style if in_callout else body_style
                target.append(Paragraph(render_inline(b[1]), style))
            elif kind == "blockquote":
                for line in b[1].split("\n"):
                    if line.strip():
                        target.append(Paragraph(render_inline(line.strip()), quote_style))
            elif kind == "code":
                target.append(Preformatted(b[2], code_style))
            elif kind == "hr":
                pass
            elif kind in ("ul", "ol"):
                bullet = "•" if kind == "ul" else None
                for idx, item in enumerate(b[1]):
                    prefix = f"{bullet} " if bullet else f"{idx + 1}. "
                    target.append(Paragraph(prefix + render_inline(item), li_style))
            elif kind == "table":
                header, rows = b[1], b[2]
                data = [[Paragraph(f"<b>{render_inline(c)}</b>", body_style) for c in header]]
                for r in rows:
                    data.append([Paragraph(render_inline(c), body_style) for c in r])
                t = Table(data, hAlign="LEFT", colWidths=None)
                t.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor(INDIGO)),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor(GRAY)),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 6),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor(LIGHT)]),
                ]))
                target.append(t)
                target.append(Spacer(1, 8))
            elif kind == "diagram":
                path = DIAGRAM_FILES.get(b[1])
                if path and path.exists():
                    from PIL import Image as PILImage
                    iw, ih = PILImage.open(path).size
                    max_w = PAGE_W - 2 * MARGIN
                    max_h = 5.2 * inch
                    scale = min(max_w / iw, max_h / ih, 1.0) if iw and ih else 1.0
                    img = RLImage(str(path), width=iw * scale, height=ih * scale)
                    img.hAlign = "CENTER"
                    target.append(Spacer(1, 6))
                    target.append(img)
                    target.append(Spacer(1, 4))
            elif kind == "callout":
                ctype = b[1]
                inner = b[2]
                color = CALLOUT_COLORS.get(ctype, GRAY)
                cell_content = [Paragraph(ctype.upper(), ParagraphStyle(
                    f"CLabel{ctype}", parent=callout_label_style,
                    textColor=colors.white, backColor=None))]
                render_blocks_into(inner, cell_content, in_callout=True)
                label_bar = Table([[Paragraph(f'<font color="white"><b>{ctype.upper()}</b></font>', body_style)]],
                                   colWidths=[PAGE_W - 2 * MARGIN - 16])
                label_bar.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(color)),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]))
                body_content = []
                render_blocks_into(inner, body_content, in_callout=True)
                body_cell = Table([[c] for c in body_content] or [[""]],
                                   colWidths=[PAGE_W - 2 * MARGIN - 16])
                body_cell.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor(LIGHT)),
                    ("LEFTPADDING", (0, 0), (-1, -1), 10),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                    ("TOPPADDING", (0, 0), (-1, -1), 6),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                ]))
                wrapper = Table([[label_bar], [body_cell]], colWidths=[PAGE_W - 2 * MARGIN - 16])
                wrapper.setStyle(TableStyle([
                    ("LINEBEFORE", (0, 0), (0, -1), 4, colors.HexColor(color)),
                    ("TOPPADDING", (0, 0), (-1, -1), 0),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                    ("LEFTPADDING", (0, 0), (-1, -1), 0),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 0),
                ]))
                target.append(Spacer(1, 6))
                target.append(KeepTogether([wrapper]))
                target.append(Spacer(1, 10))
            i += 1

    render_blocks_into(blocks, story)

    # ── Page templates: header/footer via onPage ────────────────────────────

    def draw_chrome(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(colors.HexColor(INDIGO))
        canvas.setLineWidth(0.6)
        canvas.line(MARGIN, PAGE_H - 0.55 * inch, PAGE_W - MARGIN, PAGE_H - 0.55 * inch)
        canvas.setFont("Helvetica", 8.5)
        canvas.setFillColor(colors.HexColor(GRAY))
        canvas.drawString(MARGIN, PAGE_H - 0.48 * inch, "MetaVibing — Field Manual")
        canvas.drawRightString(PAGE_W - MARGIN, PAGE_H - 0.48 * inch, VERSION_LABEL)
        canvas.setFont("Helvetica", 9)
        canvas.drawCentredString(PAGE_W / 2, 0.55 * inch, str(canvas.getPageNumber()))
        canvas.restoreState()

    def draw_nothing(canvas, doc):
        pass

    doc = BaseDocTemplate(str(out_path), pagesize=LETTER,
                           leftMargin=MARGIN, rightMargin=MARGIN,
                           topMargin=MARGIN, bottomMargin=MARGIN,
                           title=TITLE, author="MetaVibing")
    cover_frame = Frame(MARGIN, MARGIN, PAGE_W - 2 * MARGIN, PAGE_H - 2 * MARGIN, id="cover")
    body_frame = Frame(MARGIN, MARGIN, PAGE_W - 2 * MARGIN, PAGE_H - 2 * MARGIN, id="body")
    doc.addPageTemplates([
        PageTemplate(id="Cover", frames=[cover_frame], onPage=draw_nothing),
        PageTemplate(id="Body", frames=[body_frame], onPage=draw_chrome),
    ])
    doc.build(story)
    print("wrote", out_path)


# ═════════════════════════════════════════════════════════════════════════
# DOCX (python-docx)
# ═════════════════════════════════════════════════════════════════════════

def build_docx(blocks, out_path):
    import docx
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hexc(h):
        h = h.lstrip("#")
        return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))

    def add_page_number_field(paragraph):
        run = paragraph.add_run()
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    doc = Document()
    doc.core_properties.title = TITLE
    doc.core_properties.author = "MetaVibing"

    section = doc.sections[0]
    section.left_margin = section.right_margin = Inches(0.9)
    section.top_margin = section.bottom_margin = Inches(0.9)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = "MetaVibing — Field Manual"
    hp.runs[0].font.size = Pt(8.5)
    hp.runs[0].font.color.rgb = hexc(GRAY)
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = hexc(DARK)

    # ── Cover ──────────────────────────────────────────────────────────────
    for _ in range(4):
        doc.add_paragraph()
    logo_path = ASSETS / "logo.png"
    if logo_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo_path), width=Inches(1.3))
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run(TITLE)
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = hexc(INDIGO)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run(SUBTITLE)
    r.font.size = Pt(14)
    r.font.color.rgb = hexc(DARK)

    ver_p = doc.add_paragraph()
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ver_p.add_run(VERSION_LABEL)
    r.font.size = Pt(11)
    r.font.bold = True
    r.font.color.rgb = hexc(TEAL)

    doc.add_page_break()

    def shade_paragraph(paragraph, hex_color):
        pPr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), hex_color.lstrip("#"))
        pPr.append(shd)

    def add_left_border(paragraph, hex_color):
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "36")
        left.set(qn("w:space"), "6")
        left.set(qn("w:color"), hex_color.lstrip("#"))
        pBdr.append(left)
        pPr.append(pBdr)

    def add_inline_runs(paragraph, text, base_size=10.5, code_color=VIOLET):
        """Very small inline-markdown-to-runs converter for docx (bold/italic/code)."""
        pos = 0
        pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*)")
        for m in pattern.finditer(text):
            if m.start() > pos:
                paragraph.add_run(text[pos:m.start()])
            token = m.group(0)
            if token.startswith("**"):
                r = paragraph.add_run(token[2:-2])
                r.bold = True
            elif token.startswith("`"):
                r = paragraph.add_run(token[1:-1])
                r.font.name = "Consolas"
                r.font.color.rgb = hexc(code_color)
            else:
                r = paragraph.add_run(token[1:-1])
                r.italic = True
            pos = m.end()
        if pos < len(text):
            paragraph.add_run(text[pos:])

    def render_blocks(blocks, in_callout_color=None):
        for b in blocks:
            kind = b[0]
            if kind == "h2":
                text = b[1]
                if PART_RE.match(text):
                    doc.add_page_break()
                    p = doc.add_heading(level=1)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = p.runs[0] if p.runs else p.add_run(text)
                    r.text = text
                    r.font.color.rgb = hexc(INDIGO)
                    r.font.size = Pt(26)
                    doc.add_page_break()
                else:
                    h = doc.add_heading(level=2)
                    add_inline_runs(h, strip_inline(text))
                    for run in h.runs:
                        run.font.color.rgb = hexc(INDIGO)
            elif kind == "h3":
                h = doc.add_heading(level=3)
                add_inline_runs(h, strip_inline(b[1]))
                for run in h.runs:
                    run.font.color.rgb = hexc(TEAL)
            elif kind in ("h4", "h5", "h6"):
                h = doc.add_heading(level=4)
                add_inline_runs(h, strip_inline(b[1]))
            elif kind == "p":
                p = doc.add_paragraph()
                add_inline_runs(p, b[1])
                if in_callout_color:
                    shade_paragraph(p, LIGHT)
            elif kind == "blockquote":
                for line in b[1].split("\n"):
                    if line.strip():
                        p = doc.add_paragraph(style="Intense Quote")
                        add_inline_runs(p, line.strip())
            elif kind == "code":
                p = doc.add_paragraph()
                shade_paragraph(p, LIGHT)
                for j, line in enumerate(b[2].split("\n")):
                    r = p.add_run(("\n" if j else "") + line)
                    r.font.name = "Consolas"
                    r.font.size = Pt(8.5)
                    r.font.color.rgb = hexc(DARK)
            elif kind == "hr":
                pass
            elif kind in ("ul", "ol"):
                style = "List Bullet" if kind == "ul" else "List Number"
                for item in b[1]:
                    p = doc.add_paragraph(style=style)
                    add_inline_runs(p, item)
            elif kind == "table":
                header, rows = b[1], b[2]
                t = doc.add_table(rows=1, cols=len(header))
                t.style = "Light Grid Accent 1"
                for c, val in zip(t.rows[0].cells, header):
                    c.text = strip_inline(val)
                for row in rows:
                    cells = t.add_row().cells
                    for c, val in zip(cells, row):
                        c.text = strip_inline(val)
                doc.add_paragraph()
            elif kind == "diagram":
                path = DIAGRAM_FILES.get(b[1])
                if path and path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(path), width=Inches(5.6))
            elif kind == "callout":
                ctype, inner = b[1], b[2]
                color = CALLOUT_COLORS.get(ctype, GRAY)
                label_p = doc.add_paragraph()
                shade_paragraph(label_p, color)
                add_left_border(label_p, color)
                r = label_p.add_run(ctype.upper())
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                for ib in inner:
                    if ib[0] == "p":
                        p = doc.add_paragraph()
                        shade_paragraph(p, LIGHT)
                        add_left_border(p, color)
                        add_inline_runs(p, ib[1])
                    elif ib[0] == "code":
                        p = doc.add_paragraph()
                        shade_paragraph(p, LIGHT)
                        add_left_border(p, color)
                        for j, line in enumerate(ib[2].split("\n")):
                            r = p.add_run(("\n" if j else "") + line)
                            r.font.name = "Consolas"
                            r.font.size = Pt(8.5)
                    elif ib[0] == "blockquote":
                        p = doc.add_paragraph()
                        shade_paragraph(p, LIGHT)
                        add_left_border(p, color)
                        add_inline_runs(p, ib[1].replace("\n", " "))
                    else:
                        render_blocks([ib], in_callout_color=color)
                doc.add_paragraph()

    render_blocks(blocks)
    doc.save(str(out_path))
    print("wrote", out_path)


if __name__ == "__main__":
    blocks = load_blocks()
    build_pdf(blocks, DIST / "MetaVibing-Field-Manual-v0.1.pdf")
    build_docx(blocks, DIST / "MetaVibing-Field-Manual-v0.1.docx")
